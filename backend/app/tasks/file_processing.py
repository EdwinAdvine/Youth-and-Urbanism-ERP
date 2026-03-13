"""Drive file processing tasks: text extraction, embedding, AI analysis.

Celery tasks that run asynchronously after file upload:
1. extract_file_content — extract text from PDF/DOCX/XLSX/images (OCR)
2. generate_file_embedding — embed extracted text + store in pgvector
3. ai_analyze_file — summarize, extract entities, suggest tags, classify sensitivity
4. purge_expired_trash — daily cleanup of expired trash items
"""
from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
from datetime import datetime, timezone

from celery import shared_task

from app.core.config import settings

logger = logging.getLogger(__name__)

# Supported MIME types for text extraction
EXTRACTABLE_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # docx
    "application/msword",  # doc (limited)
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # xlsx
    "application/vnd.ms-excel",  # xls (limited)
    "text/plain",
    "text/csv",
    "text/markdown",
    "text/html",
    "application/json",
    "application/xml",
    "text/xml",
}

# Image types for OCR
OCR_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/tiff",
    "image/bmp",
    "image/webp",
}

# Max text to extract (prevent huge memory usage)
MAX_TEXT_LENGTH = 500_000  # ~500K chars


def _extract_text_from_pdf(file_data: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            for page in pdf.pages[:200]:  # limit to 200 pages
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
                if sum(len(t) for t in text_parts) > MAX_TEXT_LENGTH:
                    break
        return "\n\n".join(text_parts)[:MAX_TEXT_LENGTH]
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        return ""


def _extract_text_from_docx(file_data: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_data))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
            if sum(len(t) for t in text_parts) > MAX_TEXT_LENGTH:
                break
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        return "\n".join(text_parts)[:MAX_TEXT_LENGTH]
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""


def _extract_text_from_xlsx(file_data: bytes) -> str:
    """Extract text from XLSX using openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames[:20]:  # limit sheets
            ws = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    text_parts.append(row_text)
                row_count += 1
                if row_count > 5000:
                    break
            if sum(len(t) for t in text_parts) > MAX_TEXT_LENGTH:
                break
        wb.close()
        return "\n".join(text_parts)[:MAX_TEXT_LENGTH]
    except Exception as exc:
        logger.warning("XLSX extraction failed: %s", exc)
        return ""


def _extract_text_from_image(file_data: bytes) -> str:
    """Extract text from image using pytesseract OCR."""
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(file_data))
        # Convert to RGB if needed
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")
        text = pytesseract.image_to_string(image, lang="eng")
        return text[:MAX_TEXT_LENGTH]
    except ImportError:
        logger.warning("pytesseract not installed — skipping OCR")
        return ""
    except Exception as exc:
        logger.warning("OCR extraction failed: %s", exc)
        return ""


def _extract_text_from_plain(file_data: bytes) -> str:
    """Extract text from plain text / CSV / markdown / JSON / XML files."""
    try:
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_data.decode(encoding)[:MAX_TEXT_LENGTH]
            except UnicodeDecodeError:
                continue
        return ""
    except Exception:
        return ""


def extract_text(file_data: bytes, content_type: str, filename: str) -> str:
    """Route to the correct extractor based on content type."""
    ct = content_type.lower()

    if ct == "application/pdf":
        return _extract_text_from_pdf(file_data)
    elif ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_text_from_docx(file_data)
    elif ct in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_text_from_xlsx(file_data)
    elif ct in OCR_TYPES:
        return _extract_text_from_image(file_data)
    elif ct.startswith("text/") or ct in ("application/json", "application/xml"):
        return _extract_text_from_plain(file_data)
    else:
        return ""


def _download_file_from_minio(minio_key: str) -> bytes:
    """Download file content from MinIO."""
    from app.integrations.minio_client import BUCKET_NAME, _ensure_bucket, _get_client

    client = _get_client()
    _ensure_bucket(client)
    response = client.get_object(Bucket=BUCKET_NAME, Key=minio_key)
    return response["Body"].read()


# ── Celery Tasks ──────────────────────────────────────────────────────────────


@shared_task(name="tasks.extract_file_content", bind=True, max_retries=2)
def extract_file_content(self, file_id: str) -> dict:
    """Extract text content from a Drive file and store it.

    Triggered automatically after file upload for supported content types.
    """
    async def _process():
        from app.core.database import AsyncSessionLocal
        from app.models.drive import DriveFile

        async with AsyncSessionLocal() as db:
            from sqlalchemy import select
            stmt = select(DriveFile).where(DriveFile.id == file_id)
            result = await db.execute(stmt)
            drive_file = result.scalar_one_or_none()

            if not drive_file:
                return {"status": "error", "error": "File not found"}

            # Check if content type is extractable
            ct = drive_file.content_type.lower()
            if ct not in EXTRACTABLE_TYPES and ct not in OCR_TYPES:
                return {"status": "skipped", "reason": f"Unsupported type: {ct}"}

            # Download from MinIO
            try:
                file_data = _download_file_from_minio(drive_file.minio_key)
            except Exception as exc:
                return {"status": "error", "error": f"Download failed: {exc}"}

            # Extract text
            text = extract_text(file_data, drive_file.content_type, drive_file.name)
            if not text.strip():
                drive_file.ai_processed = True
                await db.commit()
                return {"status": "empty", "reason": "No text extracted"}

            # Compute content hash
            content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()

            # Store extracted text
            drive_file.file_content_text = text
            drive_file.content_hash = content_hash
            await db.commit()

            logger.info(
                "Extracted %d chars from file %s (%s)",
                len(text), drive_file.name, drive_file.content_type,
            )

            # Chain: trigger embedding generation
            generate_file_embedding.delay(file_id)
            # Chain: trigger AI analysis
            ai_analyze_file.delay(file_id)

            return {
                "status": "success",
                "file_id": file_id,
                "chars_extracted": len(text),
                "content_hash": content_hash,
            }

    try:
        return asyncio.run(_process())
    except Exception as exc:
        logger.exception("extract_file_content failed for %s", file_id)
        raise self.retry(exc=exc, countdown=30)


@shared_task(name="tasks.generate_file_embedding", bind=True, max_retries=2)
def generate_file_embedding(self, file_id: str) -> dict:
    """Generate vector embedding for a Drive file's extracted text.

    Stores both the file-level embedding (on DriveFile.content_embedding)
    and chunked embeddings (in DocumentEmbedding table for RAG search).
    """
    async def _embed():
        from app.core.database import AsyncSessionLocal
        from app.models.drive import DriveFile
        from app.services.embedding import embedding_svc

        async with AsyncSessionLocal() as db:
            from sqlalchemy import select
            stmt = select(DriveFile).where(DriveFile.id == file_id)
            result = await db.execute(stmt)
            drive_file = result.scalar_one_or_none()

            if not drive_file or not drive_file.file_content_text:
                return {"status": "skipped", "reason": "No text to embed"}

            text = drive_file.file_content_text

            try:
                # 1. Generate file-level embedding (first 8000 chars for summary embedding)
                summary_text = text[:8000]
                file_embedding = await embedding_svc.embed_text(summary_text)
                drive_file.content_embedding = file_embedding

                # 2. Chunk and embed for RAG search (using existing DocumentEmbedding)
                import uuid as uuid_mod
                chunk_count = await embedding_svc.chunk_and_embed(
                    source_type="drive_file",
                    source_id=uuid_mod.UUID(file_id),
                    text=text,
                    db=db,
                    metadata={"filename": drive_file.name, "content_type": drive_file.content_type},
                )

                await db.commit()

                logger.info(
                    "Generated embedding for file %s: %d chunks",
                    drive_file.name, chunk_count,
                )
                return {
                    "status": "success",
                    "file_id": file_id,
                    "chunks": chunk_count,
                    "embedding_dim": len(file_embedding),
                }
            except Exception as exc:
                logger.warning("Embedding generation failed for %s: %s", file_id, exc)
                return {"status": "error", "error": str(exc)}

    try:
        return asyncio.run(_embed())
    except Exception as exc:
        logger.exception("generate_file_embedding failed for %s", file_id)
        raise self.retry(exc=exc, countdown=60)


async def _ai_chat(messages, model=None):
    """Route AI chat to the configured provider (openai, grok, or anthropic)."""
    from openai import AsyncOpenAI
    from app.core.config import settings
    provider = settings.AI_PROVIDER
    if provider == "anthropic":
        import anthropic
        client = anthropic.AsyncAnthropic(api_key=settings.AI_API_KEY)
        system_parts = [m["content"] for m in messages if m["role"] == "system"]
        non_system = [m for m in messages if m["role"] != "system"]
        kwargs = {"model": model or settings.AI_MODEL, "max_tokens": 4096, "messages": non_system}
        if system_parts:
            kwargs["system"] = "\n\n".join(system_parts)
        resp = await client.messages.create(**kwargs)
        return resp.content[0].text
    else:
        client = AsyncOpenAI(api_key=settings.AI_API_KEY, base_url=settings.AI_BASE_URL)
        resp = await client.chat.completions.create(model=model or settings.AI_MODEL, messages=messages)
        return resp.choices[0].message.content or ""


@shared_task(name="tasks.ai_analyze_file", bind=True, max_retries=2)
def ai_analyze_file(self, file_id: str) -> dict:
    """AI analysis of a Drive file: summarize, extract entities, suggest tags, classify."""
    async def _analyze():
        from app.core.database import AsyncSessionLocal
        from app.models.drive import DriveFile, FileAIMetadata

        async with AsyncSessionLocal() as db:
            from sqlalchemy import select
            stmt = select(DriveFile).where(DriveFile.id == file_id)
            result = await db.execute(stmt)
            drive_file = result.scalar_one_or_none()

            if not drive_file or not drive_file.file_content_text:
                return {"status": "skipped", "reason": "No text to analyze"}

            text = drive_file.file_content_text[:10000]  # first 10K chars for analysis
            word_count = len(drive_file.file_content_text.split())

            # Build the analysis prompt
            prompt = f"""Analyze this document and provide a JSON response with the following fields:

1. "summary": A concise 2-3 sentence summary of the document.
2. "entities": An object with keys "people", "organizations", "dates", "amounts", "locations" — each an array of strings found in the document.
3. "suggested_tags": An array of 3-7 relevant tags for categorizing this document.
4. "sensitivity": One of "public", "internal", "confidential", "highly_confidential" — classify based on content.
5. "language": The primary language of the document (ISO 639-1 code, e.g., "en").
6. "document_type": What kind of document this is (e.g., "invoice", "contract", "report", "memo", "resume", "spreadsheet").
7. "module_suggestions": An array of objects with "module" and "action" keys suggesting ERP cross-links (e.g., {{"module": "finance", "action": "Create expense from invoice"}}).

Document filename: {drive_file.name}
Content type: {drive_file.content_type}

Document content:
---
{text}
---

Respond with ONLY valid JSON, no markdown formatting."""

            try:
                ai_text = await _ai_chat([
                    {"role": "system", "content": "You are a document analysis assistant. Always respond with valid JSON only."},
                    {"role": "user", "content": prompt},
                ])

                # Parse the JSON response
                try:
                    analysis = json.loads(ai_text)
                except json.JSONDecodeError:
                    # Try to extract JSON from the response
                    import re
                    json_match = re.search(r'\{[\s\S]*\}', ai_text)
                    if json_match:
                        analysis = json.loads(json_match.group())
                    else:
                        analysis = {}

                # Upsert FileAIMetadata
                from sqlalchemy import select as sa_select
                existing = await db.execute(
                    sa_select(FileAIMetadata).where(FileAIMetadata.file_id == file_id)
                )
                ai_meta = existing.scalar_one_or_none()

                if not ai_meta:
                    ai_meta = FileAIMetadata(file_id=drive_file.id)
                    db.add(ai_meta)

                ai_meta.summary = analysis.get("summary", "")
                ai_meta.entities_json = analysis.get("entities", {})
                ai_meta.suggested_tags = analysis.get("suggested_tags", [])
                ai_meta.sensitivity_level = analysis.get("sensitivity", "internal")
                ai_meta.language = analysis.get("language", "en")
                ai_meta.word_count = word_count
                ai_meta.processed_at = datetime.now(timezone.utc)
                ai_meta.module_suggestions = analysis.get("module_suggestions", [])

                # Update file sensitivity level
                drive_file.sensitivity_level = analysis.get("sensitivity", "internal")
                drive_file.ai_processed = True

                await db.commit()

                logger.info(
                    "AI analysis complete for file %s: %s",
                    drive_file.name, analysis.get("document_type", "unknown"),
                )
                return {
                    "status": "success",
                    "file_id": file_id,
                    "summary_length": len(analysis.get("summary", "")),
                    "entities_count": sum(
                        len(v) for v in analysis.get("entities", {}).values() if isinstance(v, list)
                    ),
                    "tags": analysis.get("suggested_tags", []),
                    "sensitivity": analysis.get("sensitivity", "unknown"),
                }

            except Exception as exc:
                # Store error but don't fail permanently
                from sqlalchemy import select as sa_select
                existing = await db.execute(
                    sa_select(FileAIMetadata).where(FileAIMetadata.file_id == file_id)
                )
                ai_meta = existing.scalar_one_or_none()
                if not ai_meta:
                    ai_meta = FileAIMetadata(file_id=drive_file.id)
                    db.add(ai_meta)
                ai_meta.processing_error = str(exc)
                ai_meta.processed_at = datetime.now(timezone.utc)
                drive_file.ai_processed = True
                await db.commit()

                logger.warning("AI analysis failed for %s: %s", file_id, exc)
                return {"status": "error", "error": str(exc)}

    try:
        return asyncio.run(_analyze())
    except Exception as exc:
        logger.exception("ai_analyze_file failed for %s", file_id)
        raise self.retry(exc=exc, countdown=60)


@shared_task(name="tasks.purge_expired_trash")
def purge_expired_trash() -> dict:
    """Daily task: permanently delete files whose auto_purge_at has passed."""
    async def _purge():
        from sqlalchemy import delete, select
        from app.core.database import AsyncSessionLocal
        from app.integrations.minio_client import delete_file as minio_delete
        from app.models.drive import DriveFile, TrashBin

        async with AsyncSessionLocal() as db:
            now = datetime.now(timezone.utc)

            # Find expired trash items
            stmt = select(TrashBin).where(
                TrashBin.auto_purge_at.isnot(None),
                TrashBin.auto_purge_at <= now,
            )
            result = await db.execute(stmt)
            expired_items = result.scalars().all()

            deleted_count = 0
            errors = []

            for item in expired_items:
                try:
                    # Get the file to delete from MinIO
                    file_stmt = select(DriveFile).where(DriveFile.id == item.file_id)
                    file_result = await db.execute(file_stmt)
                    drive_file = file_result.scalar_one_or_none()

                    if drive_file:
                        # Delete from MinIO
                        minio_delete(drive_file.minio_key)
                        # Delete file record (cascades to trash_bin, tags, comments, etc.)
                        await db.delete(drive_file)
                    else:
                        # Just delete the orphaned trash record
                        await db.delete(item)

                    deleted_count += 1
                except Exception as exc:
                    errors.append({"file_id": str(item.file_id), "error": str(exc)})
                    logger.warning("Failed to purge trash item %s: %s", item.file_id, exc)

            await db.commit()
            logger.info("Purged %d expired trash items (%d errors)", deleted_count, len(errors))
            return {
                "status": "success",
                "purged": deleted_count,
                "errors": errors,
            }

    return asyncio.run(_purge())


@shared_task(name="tasks.create_drive_snapshot")
def create_drive_snapshot(user_id: str) -> dict:
    """Create a point-in-time snapshot of a user's Drive for restore purposes."""
    async def _snapshot():
        from sqlalchemy import select
        from app.core.database import AsyncSessionLocal
        from app.models.drive import DriveFile, DriveFolder, DriveSnapshot

        async with AsyncSessionLocal() as db:
            # Gather all user's files
            files_stmt = select(DriveFile).where(DriveFile.owner_id == user_id)
            files_result = await db.execute(files_stmt)
            files = files_result.scalars().all()

            # Gather all user's folders
            folders_stmt = select(DriveFolder).where(DriveFolder.owner_id == user_id)
            folders_result = await db.execute(folders_stmt)
            folders = folders_result.scalars().all()

            metadata = {
                "files": [
                    {
                        "id": str(f.id),
                        "name": f.name,
                        "minio_key": f.minio_key,
                        "content_type": f.content_type,
                        "size": f.size,
                        "folder_id": str(f.folder_id) if f.folder_id else None,
                        "folder_path": f.folder_path,
                    }
                    for f in files
                ],
                "folders": [
                    {
                        "id": str(fo.id),
                        "name": fo.name,
                        "parent_id": str(fo.parent_id) if fo.parent_id else None,
                    }
                    for fo in folders
                ],
            }

            snapshot = DriveSnapshot(
                owner_id=user_id,
                metadata_json=metadata,
                file_count=len(files),
                total_size=sum(f.size for f in files),
            )
            db.add(snapshot)
            await db.commit()
            await db.refresh(snapshot)

            logger.info(
                "Created Drive snapshot for user %s: %d files, %d bytes",
                user_id, len(files), snapshot.total_size,
            )
            return {
                "status": "success",
                "snapshot_id": str(snapshot.id),
                "file_count": len(files),
                "total_size": snapshot.total_size,
            }

    return asyncio.run(_snapshot())
