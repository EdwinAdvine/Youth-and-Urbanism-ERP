"""Agentic Document Copilot — multi-step document creation from ERP data."""
from __future__ import annotations

import io
import uuid
from datetime import date, datetime
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession


class DocAgentService:
    """Creates complex documents by gathering data from multiple ERP modules."""

    def __init__(self, db: AsyncSession, user_id: uuid.UUID) -> None:
        self.db = db
        self.user_id = user_id

    async def _ai_chat(self, prompt: str, system_prompt: str | None = None) -> str:
        """Call AI service for text generation."""
        try:
            from app.services.ai import AIService
            svc = AIService(self.db)
            result = await svc.chat(self.user_id, prompt, system_prompt=system_prompt)
            return result.get("content", "")
        except Exception:
            return ""

    async def _gather_finance_data(self, start: date | None = None, end: date | None = None) -> dict[str, Any]:
        """Gather financial summary data."""
        from app.models.finance import Invoice, Expense

        inv_query = select(
            func.coalesce(func.sum(Invoice.total), 0).label("revenue"),
            func.count(Invoice.id).label("invoice_count"),
        ).where(Invoice.status != "draft")
        exp_query = select(func.coalesce(func.sum(Expense.amount), 0).label("expenses"))

        if start:
            inv_query = inv_query.where(Invoice.issue_date >= start)
            exp_query = exp_query.where(Expense.expense_date >= start)
        if end:
            inv_query = inv_query.where(Invoice.issue_date <= end)
            exp_query = exp_query.where(Expense.expense_date <= end)

        inv_result = await self.db.execute(inv_query)
        inv_row = inv_result.one()
        exp_result = await self.db.execute(exp_query)
        exp_total = exp_result.scalar() or 0

        return {
            "revenue": float(inv_row.revenue),
            "invoice_count": inv_row.invoice_count,
            "expenses": float(exp_total),
            "net_income": float(inv_row.revenue) - float(exp_total),
        }

    async def _gather_hr_data(self) -> dict[str, Any]:
        """Gather HR summary data."""
        from app.models.hr import Employee, Department

        headcount = (await self.db.execute(
            select(func.count(Employee.id)).where(Employee.status == "active")
        )).scalar() or 0

        dept_count = (await self.db.execute(select(func.count(Department.id)))).scalar() or 0

        return {"active_employees": headcount, "departments": dept_count}

    async def _gather_crm_data(self) -> dict[str, Any]:
        """Gather CRM summary data."""
        from app.models.crm import Deal, Opportunity

        pipeline_value = (await self.db.execute(
            select(func.coalesce(func.sum(Opportunity.expected_value), 0))
        )).scalar() or 0

        deal_count = (await self.db.execute(select(func.count(Deal.id)))).scalar() or 0
        won_deals = (await self.db.execute(
            select(func.count(Deal.id)).where(Deal.status == "won")
        )).scalar() or 0

        return {
            "pipeline_value": float(pipeline_value),
            "total_deals": deal_count,
            "won_deals": won_deals,
        }

    async def _gather_project_data(self) -> dict[str, Any]:
        """Gather project summary data."""
        from app.models.projects import Project, Task

        total_projects = (await self.db.execute(select(func.count(Project.id)))).scalar() or 0
        active_projects = (await self.db.execute(
            select(func.count(Project.id)).where(Project.status == "active")
        )).scalar() or 0
        total_tasks = (await self.db.execute(select(func.count(Task.id)))).scalar() or 0
        done_tasks = (await self.db.execute(
            select(func.count(Task.id)).where(Task.status == "done")
        )).scalar() or 0

        return {
            "total_projects": total_projects,
            "active_projects": active_projects,
            "total_tasks": total_tasks,
            "completed_tasks": done_tasks,
            "completion_rate": round((done_tasks / total_tasks) * 100, 1) if total_tasks > 0 else 0,
        }

    def _make_docx_doc(self) -> Any:
        """Create a base python-docx Document with brand defaults."""
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Open Sans"
        style.font.size = Pt(11)
        return doc

    def _add_branded_title(self, doc: Any, title: str) -> None:
        """Add a center-aligned, primary-coloured title heading."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x51, 0x45, 0x9D)

    def _doc_to_bytes(self, doc: Any) -> bytes:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Public document generation methods ───────────────────────────────────

    async def create_board_deck(self, start: date, end: date) -> tuple[bytes, str]:
        """Create a board meeting deck for the given date range.

        Returns (file_bytes, filename).
        """
        finance = await self._gather_finance_data(start, end)
        hr = await self._gather_hr_data()
        crm = await self._gather_crm_data()
        projects = await self._gather_project_data()

        data_summary = (
            f"Period: {start} to {end}\n"
            f"Revenue: ${finance['revenue']:,.2f}, Expenses: ${finance['expenses']:,.2f}, "
            f"Net Income: ${finance['net_income']:,.2f}, Invoices: {finance['invoice_count']}\n"
            f"Headcount: {hr['active_employees']}, Departments: {hr['departments']}\n"
            f"Pipeline: ${crm['pipeline_value']:,.2f}, Deals: {crm['total_deals']} (Won: {crm['won_deals']})\n"
            f"Projects: {projects['active_projects']} active, Tasks: {projects['completion_rate']}% complete"
        )

        exec_summary = await self._ai_chat(
            f"Write a concise 3-paragraph executive summary for a board meeting based on this data:\n{data_summary}",
            system_prompt="You are a business analyst writing board meeting documents. Be concise and data-driven.",
        )

        doc = self._make_docx_doc()
        self._add_branded_title(doc, f"Board Meeting Report — {start} to {end}")
        doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")
        doc.add_paragraph("")

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(exec_summary or "No AI summary available.")

        doc.add_heading("Financial Highlights", level=1)
        table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
        metrics = [
            ("Revenue", f"${finance['revenue']:,.2f}"),
            ("Expenses", f"${finance['expenses']:,.2f}"),
            ("Net Income", f"${finance['net_income']:,.2f}"),
            ("Invoices Issued", str(finance['invoice_count'])),
            ("Profit Margin", f"{(finance['net_income'] / finance['revenue'] * 100):.1f}%" if finance['revenue'] > 0 else "N/A"),
        ]
        for i, (label, value) in enumerate(metrics):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_heading("Human Resources", level=1)
        doc.add_paragraph(f"Active Employees: {hr['active_employees']}")
        doc.add_paragraph(f"Departments: {hr['departments']}")

        doc.add_heading("Sales & CRM Pipeline", level=1)
        doc.add_paragraph(f"Pipeline Value: ${crm['pipeline_value']:,.2f}")
        win_rate = round((crm['won_deals'] / crm['total_deals']) * 100, 1) if crm['total_deals'] > 0 else 0
        doc.add_paragraph(f"Deals: {crm['total_deals']} total, {crm['won_deals']} won ({win_rate}% win rate)")

        doc.add_heading("Project Portfolio", level=1)
        doc.add_paragraph(f"Active Projects: {projects['active_projects']} of {projects['total_projects']}")
        doc.add_paragraph(f"Task Completion: {projects['completion_rate']}%")

        filename = f"Board_Report_{start.isoformat()}_{end.isoformat()}.docx"
        return self._doc_to_bytes(doc), filename

    async def create_monthly_report(self, year: int, month: int, department: str | None = None) -> tuple[bytes, str]:
        """Create a monthly department report.

        Returns (file_bytes, filename).
        """
        start = date(year, month, 1)
        if month == 12:
            end = date(year + 1, 1, 1)
        else:
            end = date(year, month + 1, 1)

        finance = await self._gather_finance_data(start, end)
        hr = await self._gather_hr_data()
        projects = await self._gather_project_data()

        doc = self._make_docx_doc()
        month_name = start.strftime("%B %Y")
        self._add_branded_title(doc, f"Monthly Report — {month_name}")

        if department:
            doc.add_paragraph(f"Department: {department}")

        doc.add_heading("Financial Summary", level=1)
        doc.add_paragraph(f"Revenue: ${finance['revenue']:,.2f}")
        doc.add_paragraph(f"Expenses: ${finance['expenses']:,.2f}")
        doc.add_paragraph(f"Net: ${finance['net_income']:,.2f}")

        doc.add_heading("Team", level=1)
        doc.add_paragraph(f"Active employees: {hr['active_employees']}")
        doc.add_paragraph(f"Departments: {hr['departments']}")

        doc.add_heading("Projects", level=1)
        doc.add_paragraph(f"Active: {projects['active_projects']}")
        doc.add_paragraph(f"Task completion: {projects['completion_rate']}%")

        dept_slug = department.replace(" ", "_") if department else "all"
        filename = f"Monthly_Report_{dept_slug}_{year}_{month:02d}.docx"
        return self._doc_to_bytes(doc), filename

    async def create_proposal(self, deal_id: uuid.UUID) -> tuple[bytes, str]:
        """Create a sales proposal from CRM deal data.

        Returns (file_bytes, filename).
        """
        from app.models.crm import Deal, Contact

        deal = await self.db.get(Deal, deal_id)
        if not deal:
            raise ValueError(f"Deal {deal_id} not found")

        contact = await self.db.get(Contact, deal.contact_id) if deal.contact_id else None

        proposal_text = await self._ai_chat(
            f"Write a professional sales proposal for:\n"
            f"Deal: {deal.title}\nValue: ${float(deal.deal_value or 0):,.2f}\n"
            f"Client: {contact.name if contact else 'N/A'}\n"
            f"Company: {contact.company if contact else 'N/A'}",
            system_prompt=(
                "You are a sales professional writing compelling proposals. "
                "Include sections: Introduction, Proposed Solution, Pricing, Timeline, Terms."
            ),
        )

        doc = self._make_docx_doc()
        self._add_branded_title(doc, f"Proposal: {deal.title}")
        doc.add_paragraph(f"Prepared for: {contact.name if contact else 'N/A'}")
        doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Deal Value: ${float(deal.deal_value or 0):,.2f}")
        doc.add_paragraph("")

        for paragraph in (proposal_text or "Proposal content pending.").split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        safe_title = deal.title.replace(" ", "_")[:30] if deal.title else "proposal"
        filename = f"Proposal_{safe_title}_{date.today().isoformat()}.docx"
        return self._doc_to_bytes(doc), filename

    async def create_contract(
        self,
        template_type: str = "nda",
        deal_id: uuid.UUID | None = None,
        party_name: str | None = None,
        party_company: str | None = None,
        effective_date: date | None = None,
    ) -> tuple[bytes, str]:
        """Generate a contract/NDA document.

        Args:
            template_type: "nda", "service_agreement", "employment", or "purchase_order"
            deal_id: Optional CRM deal to pull party details from
            party_name: Counterparty name (overrides deal contact)
            party_company: Counterparty company (overrides deal contact)
            effective_date: Contract effective date (defaults to today)

        Returns (file_bytes, filename).
        """
        eff_date = effective_date or date.today()
        today_str = date.today().strftime("%B %d, %Y")
        eff_str = eff_date.strftime("%B %d, %Y")

        # Optionally pull party info from CRM deal
        if deal_id and not (party_name and party_company):
            try:
                from app.models.crm import Deal, Contact
                deal = await self.db.get(Deal, deal_id)
                if deal and deal.contact_id:
                    contact = await self.db.get(Contact, deal.contact_id)
                    if contact:
                        party_name = party_name or contact.name
                        party_company = party_company or getattr(contact, "company", None)
            except Exception:
                pass

        party_name = party_name or "Counterparty"
        party_company = party_company or party_name

        template_labels = {
            "nda": "Non-Disclosure Agreement",
            "service_agreement": "Service Agreement",
            "employment": "Employment Agreement",
            "purchase_order": "Purchase Order Agreement",
        }
        contract_title = template_labels.get(template_type, "Agreement")

        contract_body = await self._ai_chat(
            f"Draft a professional {contract_title} for:\n"
            f"Between: Urban Vibes Dynamics Ltd and {party_company} ({party_name})\n"
            f"Effective Date: {eff_str}\n"
            f"Template type: {template_type}\n"
            f"Include all standard clauses for a {template_type} contract.",
            system_prompt=(
                "You are a legal document drafter. Write clear, professional contract language. "
                "Use numbered sections. Include recitals, obligations, term/termination, and signature blocks."
            ),
        )

        doc = self._make_docx_doc()
        self._add_branded_title(doc, contract_title)
        doc.add_paragraph(f"Date: {today_str}")
        doc.add_paragraph(f"Effective Date: {eff_str}")
        doc.add_paragraph(f"Between: Urban Vibes Dynamics Ltd")
        doc.add_paragraph(f"And: {party_company} ({party_name})")
        doc.add_paragraph("")

        for paragraph in (contract_body or "Contract body pending legal review.").split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        # Signature block
        doc.add_paragraph("")
        doc.add_heading("Signatures", level=1)
        sig_table = doc.add_table(rows=4, cols=2, style="Light Grid")
        sig_table.rows[0].cells[0].text = "Urban Vibes Dynamics Ltd"
        sig_table.rows[0].cells[1].text = party_company
        sig_table.rows[1].cells[0].text = "Signature: ___________________"
        sig_table.rows[1].cells[1].text = "Signature: ___________________"
        sig_table.rows[2].cells[0].text = "Name: ___________________"
        sig_table.rows[2].cells[1].text = f"Name: {party_name}"
        sig_table.rows[3].cells[0].text = f"Date: {today_str}"
        sig_table.rows[3].cells[1].text = f"Date: {today_str}"

        safe_party = party_company.replace(" ", "_")[:20]
        filename = f"{template_type.upper()}_{safe_party}_{eff_date.isoformat()}.docx"
        return self._doc_to_bytes(doc), filename

    @staticmethod
    def available_actions() -> list[dict[str, str]]:
        return [
            {
                "id": "create_board_deck",
                "name": "Board Meeting Deck",
                "description": "Comprehensive board report with Finance, HR, CRM, and Project data",
            },
            {
                "id": "create_monthly_report",
                "name": "Monthly Report",
                "description": "Monthly department summary with key metrics",
            },
            {
                "id": "create_proposal",
                "name": "Sales Proposal",
                "description": "Professional proposal from CRM deal data",
            },
            {
                "id": "create_contract",
                "name": "Contract / NDA",
                "description": "Generate NDA, service agreement, employment or purchase order contract",
            },
        ]
