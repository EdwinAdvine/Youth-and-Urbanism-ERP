"""AI document generation service: creates .docx and .xlsx files, uploads to MinIO."""
from __future__ import annotations

import io
import logging
import uuid
from datetime import UTC, datetime
from typing import Any

import boto3
from botocore.client import Config as BotoConfig
from docx import Document as DocxDocument
from openpyxl import Workbook

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.drive import DriveFile

logger = logging.getLogger(__name__)

BUCKET_NAME = "urban-erp-drive"


def _get_minio_client():
    """Create a boto3 S3 client pointing at MinIO."""
    return boto3.client(
        "s3",
        endpoint_url=settings.MINIO_URL,
        aws_access_key_id=settings.MINIO_ACCESS_KEY,
        aws_secret_access_key=settings.MINIO_SECRET_KEY,
        config=BotoConfig(signature_version="s3v4"),
        region_name="us-east-1",
    )


def _ensure_bucket(s3_client) -> None:
    """Create the drive bucket if it does not exist."""
    try:
        s3_client.head_bucket(Bucket=BUCKET_NAME)
    except Exception:
        s3_client.create_bucket(Bucket=BUCKET_NAME)


class DocGenService:
    """Generates office documents and stores them in MinIO / DriveFile."""

    # ── DOCX generation ────────────────────────────────────────────────────
    async def generate_docx(
        self,
        title: str,
        content: str,
        user_id: uuid.UUID,
        db: AsyncSession,
    ) -> dict[str, Any]:
        """Create a .docx from *title* and *content*, upload to MinIO, and
        create a DriveFile record.

        *content* is expected to be plain text; paragraphs are split on
        double newlines.
        """
        doc = DocxDocument()
        doc.add_heading(title, level=1)
        for para in content.split("\n\n"):
            stripped = para.strip()
            if stripped:
                doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        file_bytes = buf.getvalue()

        file_name = f"{title.replace(' ', '_')}.docx"
        minio_key = f"ai-docs/{user_id}/{uuid.uuid4()}/{file_name}"

        s3 = _get_minio_client()
        _ensure_bucket(s3)
        s3.put_object(
            Bucket=BUCKET_NAME,
            Key=minio_key,
            Body=file_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        drive_file = DriveFile(
            name=file_name,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size=len(file_bytes),
            minio_key=minio_key,
            folder_path="/ai-docs",
            owner_id=user_id,
        )
        db.add(drive_file)
        await db.flush()
        await db.refresh(drive_file)

        logger.info("Generated DOCX %s for user %s", file_name, user_id)
        return {
            "file_id": str(drive_file.id),
            "file_name": file_name,
            "download_url": f"/api/v1/drive/files/{drive_file.id}/download",
        }

    # ── XLSX generation ────────────────────────────────────────────────────
    async def generate_xlsx(
        self,
        title: str,
        data: list[list[Any]],
        user_id: uuid.UUID,
        db: AsyncSession,
    ) -> dict[str, Any]:
        """Create an .xlsx from *data* (list of rows, first row = headers),
        upload to MinIO, and create a DriveFile record.
        """
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel sheet names max 31 chars

        for row in data:
            ws.append(row)

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        file_bytes = buf.getvalue()

        file_name = f"{title.replace(' ', '_')}.xlsx"
        minio_key = f"ai-docs/{user_id}/{uuid.uuid4()}/{file_name}"

        s3 = _get_minio_client()
        _ensure_bucket(s3)
        s3.put_object(
            Bucket=BUCKET_NAME,
            Key=minio_key,
            Body=file_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        drive_file = DriveFile(
            name=file_name,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            size=len(file_bytes),
            minio_key=minio_key,
            folder_path="/ai-docs",
            owner_id=user_id,
        )
        db.add(drive_file)
        await db.flush()
        await db.refresh(drive_file)

        logger.info("Generated XLSX %s for user %s", file_name, user_id)
        return {
            "file_id": str(drive_file.id),
            "file_name": file_name,
            "download_url": f"/api/v1/drive/files/{drive_file.id}/download",
        }


# Module-level singleton
docgen_svc = DocGenService()
