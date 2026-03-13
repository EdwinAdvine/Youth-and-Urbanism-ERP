"""ERP Template Engine — generates DOCX/XLSX documents from ERP data.

Supports: Invoice, Payslip, Purchase Order, Project Report, Financial Report,
CRM Pipeline Report. Each generator queries the database, builds a document
using python-docx or openpyxl, and returns bytes ready for MinIO upload.
"""
from __future__ import annotations

import io
import uuid
from datetime import date, datetime
from decimal import Decimal
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


# ── Brand constants ──────────────────────────────────────────────────────────

_PRIMARY = RGBColor(0x51, 0x45, 0x9D)
_SUCCESS = RGBColor(0x6F, 0xD9, 0x43)
_HEADER_FILL = PatternFill(start_color="51459D", end_color="51459D", fill_type="solid")
_HEADER_FONT = Font(name="Open Sans", bold=True, color="FFFFFF", size=11)
_BODY_FONT = Font(name="Open Sans", size=10)
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _fmt_currency(amount: Decimal | float | int | None, currency: str = "USD") -> str:
    if amount is None:
        return f"{currency} 0.00"
    return f"{currency} {float(amount):,.2f}"


def _fmt_date(d: date | datetime | None) -> str:
    if d is None:
        return "N/A"
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%B %d, %Y")


def _add_branded_heading(doc: DocxDocument, text: str) -> None:
    """Add a branded heading to a DOCX document."""
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = _PRIMARY
        run.font.name = "Open Sans"


def _style_xlsx_header(ws, row: int, col_count: int) -> None:
    """Style a header row in a worksheet."""
    for col in range(1, col_count + 1):
        cell = ws.cell(row=row, column=col)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = _THIN_BORDER


# ── Template Engine ──────────────────────────────────────────────────────────


class ERPTemplateEngine:
    """Generates ERP documents from live database data."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    # ── Invoice DOCX ─────────────────────────────────────────────────────

    async def generate_invoice_docx(self, invoice_id: uuid.UUID) -> tuple[bytes, str]:
        """Generate a formatted invoice DOCX. Returns (bytes, filename)."""
        from app.models.finance import Invoice  # noqa: PLC0415

        invoice = await self.db.get(Invoice, invoice_id)
        if not invoice:
            raise ValueError(f"Invoice {invoice_id} not found")

        doc = DocxDocument()
        _add_branded_heading(doc, "INVOICE")

        # Invoice details table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        details = [
            ("Invoice Number:", invoice.invoice_number),
            ("Date:", _fmt_date(invoice.issue_date)),
            ("Due Date:", _fmt_date(invoice.due_date)),
            ("Status:", invoice.status.upper()),
        ]
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            for cell in table.row_cells(i):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Open Sans"
                        run.font.size = Pt(10)

        doc.add_paragraph("")

        # Bill To
        p = doc.add_paragraph()
        run = p.add_run("Bill To: ")
        run.bold = True
        run.font.name = "Open Sans"
        run = p.add_run(invoice.customer_name or "—")
        run.font.name = "Open Sans"
        if invoice.customer_email:
            p2 = doc.add_paragraph(invoice.customer_email)
            for run in p2.runs:
                run.font.name = "Open Sans"
                run.font.size = Pt(9)

        doc.add_paragraph("")

        # Line items
        items = invoice.items or []
        if items:
            items_table = doc.add_table(rows=1 + len(items), cols=4)
            items_table.style = "Table Grid"
            headers = ["Description", "Qty", "Unit Price", "Amount"]
            for j, h in enumerate(headers):
                cell = items_table.cell(0, j)
                cell.text = h
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.name = "Open Sans"
                        run.font.size = Pt(9)

            for idx, item in enumerate(items, start=1):
                desc = item.get("description", item.get("name", ""))
                qty = item.get("quantity", 1)
                price = item.get("unit_price", item.get("price", 0))
                amount = item.get("amount", float(qty) * float(price))
                items_table.cell(idx, 0).text = str(desc)
                items_table.cell(idx, 1).text = str(qty)
                items_table.cell(idx, 2).text = _fmt_currency(price, invoice.currency)
                items_table.cell(idx, 3).text = _fmt_currency(amount, invoice.currency)

        doc.add_paragraph("")

        # Totals
        totals = doc.add_table(rows=3, cols=2)
        totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
        totals_data = [
            ("Subtotal:", _fmt_currency(invoice.subtotal, invoice.currency)),
            ("Tax:", _fmt_currency(invoice.tax_amount, invoice.currency)),
            ("Total:", _fmt_currency(invoice.total, invoice.currency)),
        ]
        for i, (label, value) in enumerate(totals_data):
            totals.cell(i, 0).text = label
            totals.cell(i, 1).text = value
            for cell in totals.row_cells(i):
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Open Sans"
                        run.font.size = Pt(10)
                        if i == 2:
                            run.bold = True

        if invoice.notes:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Notes: ")
            run.bold = True
            run.font.name = "Open Sans"
            run = p.add_run(invoice.notes)
            run.font.name = "Open Sans"

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph("Generated by Urban Vibes Dynamics")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        filename = f"Invoice_{invoice.invoice_number}.docx"
        return buf.getvalue(), filename

    # ── Payslip DOCX ─────────────────────────────────────────────────────

    async def generate_payslip_docx(
        self, employee_id: uuid.UUID, period_start: date, period_end: date
    ) -> tuple[bytes, str]:
        """Generate a payslip DOCX. Returns (bytes, filename)."""
        from app.models.hr import Employee, Payslip  # noqa: PLC0415

        employee = await self.db.get(Employee, employee_id)
        if not employee:
            raise ValueError(f"Employee {employee_id} not found")

        # Find payslip for the period
        result = await self.db.execute(
            select(Payslip).where(
                Payslip.employee_id == employee_id,
                Payslip.period_start == period_start,
                Payslip.period_end == period_end,
            ).limit(1)
        )
        payslip = result.scalar_one_or_none()

        # Load user for name
        from app.models.user import User  # noqa: PLC0415
        user = await self.db.get(User, employee.user_id)
        emp_name = f"{user.first_name or ''} {user.last_name or ''}".strip() if user else f"EMP-{employee.employee_number}"

        doc = DocxDocument()
        _add_branded_heading(doc, "PAYSLIP")

        # Employee info
        info_table = doc.add_table(rows=4, cols=2)
        info_data = [
            ("Employee:", emp_name),
            ("Employee #:", employee.employee_number),
            ("Job Title:", employee.job_title or "—"),
            ("Period:", f"{_fmt_date(period_start)} – {_fmt_date(period_end)}"),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)

        doc.add_paragraph("")

        if payslip:
            # Earnings & Deductions
            pay_table = doc.add_table(rows=4, cols=2)
            pay_table.style = "Table Grid"
            pay_data = [
                ("Gross Pay:", _fmt_currency(payslip.gross_pay, employee.currency)),
                ("Deductions:", _fmt_currency(payslip.deductions_total, employee.currency)),
                ("Net Pay:", _fmt_currency(payslip.net_pay, employee.currency)),
                ("Status:", payslip.status.upper()),
            ]
            for i, (label, value) in enumerate(pay_data):
                pay_table.cell(i, 0).text = label
                pay_table.cell(i, 1).text = value
                for cell in pay_table.row_cells(i):
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Open Sans"
                            if i == 2:
                                run.bold = True
        else:
            doc.add_paragraph("No payslip found for this period.")

        doc.add_paragraph("")
        footer = doc.add_paragraph("Generated by Urban Vibes Dynamics — Confidential")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        filename = f"Payslip_{employee.employee_number}_{period_start.isoformat()}.docx"
        return buf.getvalue(), filename

    # ── Purchase Order DOCX ──────────────────────────────────────────────

    async def generate_purchase_order_docx(self, requisition_id: uuid.UUID) -> tuple[bytes, str]:
        """Generate a PO document from a procurement requisition."""
        from app.models.supplychain import ProcurementRequisition, RequisitionLine  # noqa: PLC0415

        req = await self.db.get(ProcurementRequisition, requisition_id)
        if not req:
            raise ValueError(f"Requisition {requisition_id} not found")

        # Load lines
        result = await self.db.execute(
            select(RequisitionLine).where(RequisitionLine.requisition_id == requisition_id)
        )
        lines = result.scalars().all()

        doc = DocxDocument()
        _add_branded_heading(doc, "PURCHASE ORDER")

        # PO details
        table = doc.add_table(rows=4, cols=2)
        details = [
            ("PO Number:", req.requisition_number),
            ("Title:", req.title),
            ("Priority:", req.priority.upper()),
            ("Required By:", _fmt_date(req.required_by_date)),
        ]
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)

        doc.add_paragraph("")

        # Items table
        if lines:
            items_table = doc.add_table(rows=1 + len(lines), cols=3)
            items_table.style = "Table Grid"
            for j, h in enumerate(["Item", "Qty", "Est. Unit Price"]):
                cell = items_table.cell(0, j)
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.name = "Open Sans"

            for idx, line in enumerate(lines, start=1):
                items_table.cell(idx, 0).text = str(line.item_id)[:8] + "…"
                items_table.cell(idx, 1).text = str(line.quantity)
                items_table.cell(idx, 2).text = _fmt_currency(line.estimated_unit_price)

        doc.add_paragraph("")

        # Total
        p = doc.add_paragraph()
        run = p.add_run(f"Total Estimated: {_fmt_currency(req.total_estimated)}")
        run.bold = True
        run.font.name = "Open Sans"
        run.font.size = Pt(12)

        if req.notes:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Notes: ")
            run.bold = True
            run = p.add_run(req.notes)

        doc.add_paragraph("")
        footer = doc.add_paragraph("Generated by Urban Vibes Dynamics")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        filename = f"PO_{req.requisition_number}.docx"
        return buf.getvalue(), filename

    # ── Project Report DOCX ──────────────────────────────────────────────

    async def generate_project_report_docx(self, project_id: uuid.UUID) -> tuple[bytes, str]:
        """Generate a project status report DOCX."""
        from app.models.projects import Milestone, Project, Task  # noqa: PLC0415

        project = await self.db.get(Project, project_id)
        if not project:
            raise ValueError(f"Project {project_id} not found")

        # Task stats
        result = await self.db.execute(
            select(
                Task.status,
                func.count(Task.id),
            )
            .where(Task.project_id == project_id)
            .group_by(Task.status)
        )
        task_stats = {row[0]: row[1] for row in result.all()}
        total_tasks = sum(task_stats.values())
        done_tasks = task_stats.get("done", 0)

        # Milestones
        result = await self.db.execute(
            select(Milestone)
            .where(Milestone.project_id == project_id)
            .order_by(Milestone.due_date.asc())
        )
        milestones = result.scalars().all()

        doc = DocxDocument()
        _add_branded_heading(doc, f"Project Report: {project.name}")

        # Overview
        doc.add_heading("Overview", level=2)
        overview_table = doc.add_table(rows=4, cols=2)
        overview_data = [
            ("Status:", project.status.upper()),
            ("Start Date:", _fmt_date(project.start_date)),
            ("End Date:", _fmt_date(project.end_date)),
            ("Progress:", f"{done_tasks}/{total_tasks} tasks completed ({(done_tasks / total_tasks * 100) if total_tasks else 0:.0f}%)"),
        ]
        for i, (label, value) in enumerate(overview_data):
            overview_table.cell(i, 0).text = label
            overview_table.cell(i, 1).text = str(value)

        if project.description:
            doc.add_paragraph("")
            doc.add_paragraph(project.description)

        # Task breakdown
        doc.add_paragraph("")
        doc.add_heading("Task Breakdown", level=2)
        status_table = doc.add_table(rows=1 + len(task_stats), cols=2)
        status_table.style = "Table Grid"
        status_table.cell(0, 0).text = "Status"
        status_table.cell(0, 1).text = "Count"
        for i, (s, count) in enumerate(sorted(task_stats.items()), start=1):
            status_table.cell(i, 0).text = s.replace("_", " ").title()
            status_table.cell(i, 1).text = str(count)

        # Milestones
        if milestones:
            doc.add_paragraph("")
            doc.add_heading("Milestones", level=2)
            ms_table = doc.add_table(rows=1 + len(milestones), cols=3)
            ms_table.style = "Table Grid"
            for j, h in enumerate(["Milestone", "Due Date", "Status"]):
                ms_table.cell(0, j).text = h
            for i, ms in enumerate(milestones, start=1):
                ms_table.cell(i, 0).text = ms.title
                ms_table.cell(i, 1).text = _fmt_date(ms.due_date)
                ms_table.cell(i, 2).text = "Completed" if ms.is_completed else "Open"

        doc.add_paragraph("")
        footer = doc.add_paragraph(f"Generated by Urban Vibes Dynamics — {_fmt_date(date.today())}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        safe_name = project.name.replace(" ", "_")[:30]
        filename = f"Project_Report_{safe_name}.docx"
        return buf.getvalue(), filename

    # ── Financial Report XLSX ────────────────────────────────────────────

    async def generate_financial_report_xlsx(
        self, report_type: str, start_date: date, end_date: date
    ) -> tuple[bytes, str]:
        """Generate a financial summary XLSX (revenue, expenses, or overview)."""
        from app.models.finance import Expense, Invoice, Payment  # noqa: PLC0415

        wb = Workbook()
        ws = wb.active
        ws.title = report_type.replace("_", " ").title()

        # Title row
        ws.merge_cells("A1:F1")
        title_cell = ws["A1"]
        title_cell.value = f"Financial Report: {report_type.replace('_', ' ').title()}"
        title_cell.font = Font(name="Open Sans", bold=True, size=14, color="51459D")
        title_cell.alignment = Alignment(horizontal="center")

        ws["A2"] = f"Period: {_fmt_date(start_date)} – {_fmt_date(end_date)}"
        ws["A2"].font = Font(name="Open Sans", size=10, italic=True)

        row = 4

        if report_type in ("revenue", "overview"):
            # Revenue section — invoices
            ws.cell(row=row, column=1, value="REVENUE (Invoices)")
            ws.cell(row=row, column=1).font = Font(name="Open Sans", bold=True, size=12, color="51459D")
            row += 1

            headers = ["Invoice #", "Customer", "Issue Date", "Due Date", "Status", "Total"]
            for col, h in enumerate(headers, 1):
                ws.cell(row=row, column=col, value=h)
            _style_xlsx_header(ws, row, len(headers))
            row += 1

            result = await self.db.execute(
                select(Invoice)
                .where(Invoice.issue_date >= start_date, Invoice.issue_date <= end_date)
                .order_by(Invoice.issue_date.asc())
            )
            invoices = result.scalars().all()
            total_revenue = Decimal("0")
            for inv in invoices:
                ws.cell(row=row, column=1, value=inv.invoice_number)
                ws.cell(row=row, column=2, value=inv.customer_name or "—")
                ws.cell(row=row, column=3, value=_fmt_date(inv.issue_date))
                ws.cell(row=row, column=4, value=_fmt_date(inv.due_date))
                ws.cell(row=row, column=5, value=inv.status)
                ws.cell(row=row, column=6, value=float(inv.total))
                ws.cell(row=row, column=6).number_format = "#,##0.00"
                for col in range(1, 7):
                    ws.cell(row=row, column=col).font = _BODY_FONT
                    ws.cell(row=row, column=col).border = _THIN_BORDER
                total_revenue += inv.total
                row += 1

            ws.cell(row=row, column=5, value="Total Revenue:")
            ws.cell(row=row, column=5).font = Font(name="Open Sans", bold=True)
            ws.cell(row=row, column=6, value=float(total_revenue))
            ws.cell(row=row, column=6).font = Font(name="Open Sans", bold=True)
            ws.cell(row=row, column=6).number_format = "#,##0.00"
            row += 2

        if report_type in ("expenses", "overview"):
            # Expenses section
            ws.cell(row=row, column=1, value="EXPENSES")
            ws.cell(row=row, column=1).font = Font(name="Open Sans", bold=True, size=12, color="51459D")
            row += 1

            headers = ["Description", "Category", "Date", "Status", "Amount"]
            for col, h in enumerate(headers, 1):
                ws.cell(row=row, column=col, value=h)
            _style_xlsx_header(ws, row, len(headers))
            row += 1

            result = await self.db.execute(
                select(Expense)
                .where(Expense.expense_date >= start_date, Expense.expense_date <= end_date)
                .order_by(Expense.expense_date.asc())
            )
            expenses = result.scalars().all()
            total_expenses = Decimal("0")
            for exp in expenses:
                ws.cell(row=row, column=1, value=exp.description)
                ws.cell(row=row, column=2, value=exp.category)
                ws.cell(row=row, column=3, value=_fmt_date(exp.expense_date))
                ws.cell(row=row, column=4, value=exp.status)
                ws.cell(row=row, column=5, value=float(exp.amount))
                ws.cell(row=row, column=5).number_format = "#,##0.00"
                for col in range(1, 6):
                    ws.cell(row=row, column=col).font = _BODY_FONT
                    ws.cell(row=row, column=col).border = _THIN_BORDER
                total_expenses += exp.amount
                row += 1

            ws.cell(row=row, column=4, value="Total Expenses:")
            ws.cell(row=row, column=4).font = Font(name="Open Sans", bold=True)
            ws.cell(row=row, column=5, value=float(total_expenses))
            ws.cell(row=row, column=5).font = Font(name="Open Sans", bold=True)
            ws.cell(row=row, column=5).number_format = "#,##0.00"
            row += 2

        # Auto-fit columns
        for col in range(1, 7):
            ws.column_dimensions[get_column_letter(col)].width = 18

        # Footer
        row += 1
        ws.cell(row=row, column=1, value=f"Generated by Urban Vibes Dynamics — {_fmt_date(date.today())}")
        ws.cell(row=row, column=1).font = Font(name="Open Sans", size=8, color="999999")

        buf = io.BytesIO()
        wb.save(buf)
        filename = f"Financial_Report_{report_type}_{start_date.isoformat()}.xlsx"
        return buf.getvalue(), filename

    # ── CRM Pipeline Report XLSX ─────────────────────────────────────────

    async def generate_crm_pipeline_xlsx(self, pipeline_id: uuid.UUID | None = None) -> tuple[bytes, str]:
        """Generate a CRM pipeline report XLSX."""
        from app.models.crm import Deal, Opportunity, Pipeline  # noqa: PLC0415

        wb = Workbook()
        ws = wb.active

        pipeline_name = "All Pipelines"
        if pipeline_id:
            pipeline = await self.db.get(Pipeline, pipeline_id)
            if pipeline:
                pipeline_name = pipeline.name

        ws.title = "Pipeline Report"

        ws.merge_cells("A1:G1")
        title_cell = ws["A1"]
        title_cell.value = f"CRM Pipeline Report: {pipeline_name}"
        title_cell.font = Font(name="Open Sans", bold=True, size=14, color="51459D")
        title_cell.alignment = Alignment(horizontal="center")

        # Opportunities
        row = 3
        ws.cell(row=row, column=1, value="OPPORTUNITIES")
        ws.cell(row=row, column=1).font = Font(name="Open Sans", bold=True, size=12, color="51459D")
        row += 1

        headers = ["Title", "Stage", "Probability", "Expected Value", "Close Date", "Assigned To"]
        for col, h in enumerate(headers, 1):
            ws.cell(row=row, column=col, value=h)
        _style_xlsx_header(ws, row, len(headers))
        row += 1

        query = select(Opportunity).order_by(Opportunity.stage.asc())
        if pipeline_id:
            query = query.where(Opportunity.pipeline_id == pipeline_id)

        result = await self.db.execute(query)
        opps = result.scalars().all()
        total_value = Decimal("0")

        for opp in opps:
            ws.cell(row=row, column=1, value=opp.title)
            ws.cell(row=row, column=2, value=opp.stage)
            ws.cell(row=row, column=3, value=f"{opp.probability or 0}%")
            ws.cell(row=row, column=4, value=float(opp.expected_value or 0))
            ws.cell(row=row, column=4).number_format = "#,##0.00"
            ws.cell(row=row, column=5, value=_fmt_date(opp.expected_close_date))
            ws.cell(row=row, column=6, value=str(opp.assigned_to)[:8] + "…" if opp.assigned_to else "—")
            for col in range(1, 7):
                ws.cell(row=row, column=col).font = _BODY_FONT
                ws.cell(row=row, column=col).border = _THIN_BORDER
            total_value += opp.expected_value or Decimal("0")
            row += 1

        ws.cell(row=row, column=3, value="Total Pipeline Value:")
        ws.cell(row=row, column=3).font = Font(name="Open Sans", bold=True)
        ws.cell(row=row, column=4, value=float(total_value))
        ws.cell(row=row, column=4).font = Font(name="Open Sans", bold=True)
        ws.cell(row=row, column=4).number_format = "#,##0.00"
        row += 2

        # Deals summary
        ws.cell(row=row, column=1, value="CLOSED DEALS")
        ws.cell(row=row, column=1).font = Font(name="Open Sans", bold=True, size=12, color="51459D")
        row += 1

        deal_headers = ["Title", "Value", "Close Date", "Status"]
        for col, h in enumerate(deal_headers, 1):
            ws.cell(row=row, column=col, value=h)
        _style_xlsx_header(ws, row, len(deal_headers))
        row += 1

        result = await self.db.execute(
            select(Deal).order_by(Deal.close_date.desc()).limit(50)
        )
        deals = result.scalars().all()
        for deal in deals:
            ws.cell(row=row, column=1, value=deal.title)
            ws.cell(row=row, column=2, value=float(deal.deal_value))
            ws.cell(row=row, column=2).number_format = "#,##0.00"
            ws.cell(row=row, column=3, value=_fmt_date(deal.close_date))
            ws.cell(row=row, column=4, value=deal.status)
            for col in range(1, 5):
                ws.cell(row=row, column=col).font = _BODY_FONT
                ws.cell(row=row, column=col).border = _THIN_BORDER
            row += 1

        # Auto-fit
        for col in range(1, 8):
            ws.column_dimensions[get_column_letter(col)].width = 20

        row += 1
        ws.cell(row=row, column=1, value=f"Generated by Urban Vibes Dynamics — {_fmt_date(date.today())}")
        ws.cell(row=row, column=1).font = Font(name="Open Sans", size=8, color="999999")

        buf = io.BytesIO()
        wb.save(buf)
        safe_name = pipeline_name.replace(" ", "_")[:20]
        filename = f"CRM_Pipeline_{safe_name}.xlsx"
        return buf.getvalue(), filename

    # ── Registry of available templates ──────────────────────────────────

    @staticmethod
    def available_templates() -> list[dict[str, str]]:
        """Return the list of ERP template types available for generation."""
        return [
            {
                "id": "invoice",
                "name": "Invoice Document",
                "module": "finance",
                "doc_type": "docx",
                "description": "Professional invoice with line items, totals, and company branding",
                "required_fields": "invoice_id",
            },
            {
                "id": "payslip",
                "name": "Employee Payslip",
                "module": "hr",
                "doc_type": "docx",
                "description": "Employee pay stub with earnings, deductions, and net pay",
                "required_fields": "employee_id, period_start, period_end",
            },
            {
                "id": "purchase_order",
                "name": "Purchase Order",
                "module": "supply_chain",
                "doc_type": "docx",
                "description": "Purchase order from a procurement requisition",
                "required_fields": "requisition_id",
            },
            {
                "id": "project_report",
                "name": "Project Status Report",
                "module": "projects",
                "doc_type": "docx",
                "description": "Project overview with task stats, milestones, and progress",
                "required_fields": "project_id",
            },
            {
                "id": "financial_report",
                "name": "Financial Report",
                "module": "finance",
                "doc_type": "xlsx",
                "description": "Revenue and expense report for a date range",
                "required_fields": "report_type (revenue|expenses|overview), start_date, end_date",
            },
            {
                "id": "crm_pipeline",
                "name": "CRM Pipeline Report",
                "module": "crm",
                "doc_type": "xlsx",
                "description": "Opportunities and deals pipeline analysis",
                "required_fields": "pipeline_id (optional)",
            },
        ]
